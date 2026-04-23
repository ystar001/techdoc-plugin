"""DOCX Exporter — HTML → DOCX 변환."""

from __future__ import annotations

import re
from pathlib import Path

from rich.console import Console

console = Console()


class DOCXExporter:
    """HTML 파일을 DOCX로 변환."""

    def export(
        self,
        html_path: Path,
        output_dir: Path,
        title: str = "",
        filename: str = "document.docx",
    ) -> Path | None:
        """HTML → DOCX 변환. python-docx 미설치 시 건너뜀."""
        try:
            from docx import Document as DocxDocument
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            console.print("  [dim]DOCX 변환 건너뜀 (python-docx 필요)[/dim]")
            return None

        output_path = output_dir / filename

        try:
            html_text = html_path.read_text(encoding="utf-8")

            doc = DocxDocument()

            # 페이지 설정
            section = doc.sections[0]
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)

            # 표지
            self._add_cover(doc, title)

            # 본문 파싱
            sections = self._extract_sections(html_text)
            for sec_title, sec_html in sections:
                # 섹션 제목
                h = doc.add_heading(sec_title, level=1)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

                # 본문
                text = re.sub(r'<[^>]+>', '', sec_html)
                text = re.sub(r'\s+', ' ', text).strip()
                paragraphs = text.split('. ')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para + ('.' if not para.endswith('.') else ''))

            doc.save(str(output_path))
            console.print(f"  DOCX: {output_path}")
            return output_path

        except Exception as e:
            console.print(f"  [yellow]DOCX 변환 실패: {str(e)[:50]}[/yellow]")
            return None

    @staticmethod
    def _add_cover(doc, title: str) -> None:
        """표지 페이지 추가."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        # 제목 파싱
        main_title = title
        part_label = ""
        chapter_label = ""

        if " - " in title:
            parts = title.split(" - ", 1)
            main_title = parts[0].strip()
            rest = parts[1].strip()

            import re
            part_match = re.search(r'(Part\s*[A-Z]\.\s*[^\d]+?)(?=\d+장|$)', rest, re.IGNORECASE)
            if part_match:
                part_label = part_match.group(1).strip()
            ch_match = re.search(r'(\d+장[\.\s]*.*)', rest)
            if ch_match:
                chapter_label = ch_match.group(1).strip()

        for _ in range(5):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(main_title)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

        if part_label:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(part_label)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x9F)

        if chapter_label:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(chapter_label)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        for _ in range(4):
            doc.add_paragraph()

        from datetime import date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date.today().strftime("%Y년 %m월"))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_page_break()

    @staticmethod
    def _extract_sections(html: str) -> list[tuple[str, str]]:
        """HTML에서 섹션 추출."""
        sections = []
        parts = re.split(r'<div class="doc-section"[^>]*>', html)
        for part in parts[1:]:
            end = part.find('</div>')
            if end == -1:
                content = part
            else:
                content = part[:end]

            title_match = re.search(r'<h2[^>]*>(.*?)</h2>', content, re.DOTALL)
            if title_match:
                title = re.sub(r'<[^>]+>', '', title_match.group(1)).strip()
                sections.append((title, content))

        return sections
